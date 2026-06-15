import io
from typing import Tuple
from pypdf import PdfReader
from docx import Document



def extract_text_from_pdf(content: bytes) -> str:
    # convert the pdf bytes into a file-like object so pypdf can read it
    reader = PdfReader(io.BytesIO(content))

    # this list will store text extracted from each page
    pages = []

    # go through every page in the pdf
    for page in reader.pages:
        try:
            # extract text from the current page
            # if no text is found, add an empty string instead
            pages.append(page.extract_text() or "")
        except Exception:
            # if one page fails, skip it instead of stopping the whole process
            continue

    # join text from all pages and remove extra spaces from start and end
    return "\n".join(pages).strip()


def extract_text_from_docx(content: bytes) -> str:
    # convert the docx bytes into a file-like object so python-docx can read it
    doc = Document(io.BytesIO(content))

    # collect all normal paragraph text from the document
    parts = [p.text for p in doc.paragraphs if p.text]

    # also check tables because resumes often contain information inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # add cell text only if it is not empty
                if cell.text:
                    parts.append(cell.text)

    # join all extracted text and remove extra spaces from start and end
    return "\n".join(parts).strip()


def extract_resume_text(filename: str, content: bytes) -> Tuple[str, str]:
    """
    returns the extracted text and the detected file extension.
    supported formats are pdf, docx, and txt.
    """

    # make filename lowercase so extension checking becomes case-insensitive
    lower = filename.lower()

    # if the file is a pdf, use the pdf extraction function
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(content), "pdf"

    # if the file is a docx, use the docx extraction function
    if lower.endswith(".docx"):
        return extract_text_from_docx(content), "docx"

    # if the file is txt, directly decode the bytes into normal text
    if lower.endswith(".txt"):
        return content.decode("utf-8", errors="ignore"), "txt"

    # if file type is not supported, show a clear error
    raise ValueError("unsupported file type. use pdf, docx, or txt.")

